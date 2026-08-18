from __future__ import annotations

import logging
from pathlib import Path

from app.document.parsers.base import BaseParser, DocumentSection

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Parse Word (.docx) files using python-docx."""

    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def parse(self, file_path: str) -> list[DocumentSection]:
        try:
            from docx import Document
            from docx.table import Table as DocxTable
        except ImportError:
            logger.warning("python-docx not installed")
            return [DocumentSection(
                title=Path(file_path).stem,
                content="[Word 解析需要安装 python-docx: pip install python-docx]",
                metadata={"file": file_path, "parser": "fallback"},
            )]

        doc = Document(file_path)
        sections: list[DocumentSection] = []
        current_heading = ""

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":  # paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                if para is None:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else ""
                section_type = "heading" if "Heading" in style_name else "paragraph"

                if section_type == "heading":
                    current_heading = text

                sections.append(DocumentSection(
                    title=text if section_type == "heading" else current_heading,
                    content=text,
                    section_type=section_type,
                    page_number=None,
                    metadata={
                        "style": style_name,
                        "bold": any(run.bold for run in para.runs if run.bold),
                    },
                ))

            elif tag == "tbl":  # table
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break
                if table is None:
                    continue

                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)

                if rows:
                    content = "\n".join(" | ".join(row) for row in rows)
                    sections.append(DocumentSection(
                        title=f"表格 - {current_heading}" if current_heading else "表格",
                        content=content,
                        section_type="table",
                        page_number=None,
                        metadata={"rows": rows, "row_count": len(rows), "col_count": len(rows[0]) if rows else 0},
                    ))

        # Extract images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_bytes = rel.target_part.blob
                    sections.append(DocumentSection(
                        title=f"图片 - {current_heading}" if current_heading else "图片",
                        content="[图片]",
                        section_type="image",
                        raw_bytes=img_bytes,
                        metadata={"content_type": rel.reltype},
                    ))
                except Exception:
                    pass

        return sections if sections else [DocumentSection(content="[空文档]")]
